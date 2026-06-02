"""文档解析函数 — 从 textbook_router 抽取。

提供 DOCX/PDF 的章节解析、内容提取功能。
textbook_router 通过 from parsers.document_parsers import ... 使用。
"""
import re
import docx
import pdfplumber
from typing import List, Dict
from io import BytesIO
from logger import get_logger

logger = get_logger()


def chinese_to_num(chinese: str) -> int:
    """中文数字转阿拉伯数字（一~十）。"""
    mapping = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
               "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
    if chinese in mapping:
        return mapping[chinese]
    if len(chinese) == 2 and chinese[0] == "十":
        return 10 + mapping.get(chinese[1], 0)
    return 0


def parse_docx(content: bytes) -> List[Dict]:
    """解析 DOCX 文件，提取章节和内容。"""
    try:
        doc = docx.Document(BytesIO(content))
        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith("Heading"):
                if current_section:
                    sections.append(current_section)
                current_section = {"title": text, "content": "", "level": int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1}
            elif current_section:
                current_section["content"] += text + "\n"
            else:
                current_section = {"title": "正文", "content": text + "\n", "level": 0}

        if current_section:
            sections.append(current_section)

        return sections
    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return []


def parse_pdf(content: bytes) -> List[Dict]:
    """解析 PDF 文件，提取页面文本。"""
    try:
        sections = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    sections.append({
                        "title": f"第{i+1}页",
                        "content": text,
                        "level": 0,
                        "page": i + 1,
                    })
        return sections
    except Exception as e:
        logger.error(f"PDF 解析失败: {e}")
        return []


def parse_docx_with_chapters(content: bytes) -> List[Dict]:
    """解析 DOCX 并提取带章节结构的内容。"""
    doc = docx.Document(BytesIO(content))
    chapters = []
    current_chapter = None
    current_section = None

    chapter_pattern = re.compile(r"^第[一二三四五六七八九十]+章")
    section_pattern = re.compile(r"^第[一二三四五六七八九十]+节")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if chapter_pattern.match(text):
            if current_chapter:
                if current_section:
                    current_chapter["sections"].append(current_section)
                chapters.append(current_chapter)
            current_chapter = {"title": text, "sections": [], "content": ""}
            current_section = None
        elif section_pattern.match(text) and current_chapter:
            if current_section:
                current_chapter["sections"].append(current_section)
            current_section = {"title": text, "content": ""}
        elif current_section:
            current_section["content"] += text + "\n"
        elif current_chapter:
            current_chapter["content"] += text + "\n"

    if current_chapter:
        if current_section:
            current_chapter["sections"].append(current_section)
        chapters.append(current_chapter)

    return chapters


def parse_pdf_with_chapters(content: bytes) -> List[Dict]:
    """解析 PDF 并尝试识别章节结构。"""
    sections = parse_pdf(content)
    chapters = []
    current_chapter = None

    chapter_pattern = re.compile(r"^第[一二三四五六七八九十]+章")

    for section in sections:
        text = section["content"]
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if chapter_pattern.match(line):
                if current_chapter:
                    chapters.append(current_chapter)
                current_chapter = {"title": line, "content": "", "page": section.get("page", 0)}
            elif current_chapter:
                current_chapter["content"] += line + "\n"

    if current_chapter:
        chapters.append(current_chapter)

    if not chapters:
        return [{"title": s["title"], "content": s["content"]} for s in sections]

    return chapters
