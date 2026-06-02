"""
高考真题提取器
从 Word 文档中提取高考真题，通过 AI 结构化处理后入库
"""
import os
import json
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from docx import Document
from docx.table import Table
from openai import OpenAI
import base64
import io
import time

from logger import get_logger
from config import PROMPT_DIR

logger = get_logger()


@dataclass
class ExtractedContent:
    """提取的文档内容"""
    text: str
    tables: List[str]
    images: List[bytes]
    filename: str
    category: str  # 知识点分类


class GaokaoExtractor:
    """高考真题提取器"""

    def __init__(self, api_key: str, api_base: str = None):
        self.api_key = api_key
        self.api_base = api_base or "https://www.chataiapi.com/v1"
        self.client = OpenAI(api_key=self.api_key, base_url=self.api_base)
        self.model = "deepseek-chat"

        # 加载 prompt
        self.prompt_template = self._load_prompt()

        # API 频率控制
        self.last_request_time = 0
        self.min_request_interval = 1.0  # 1秒间隔

        logger.info(f"高考真题提取器初始化完成")
        logger.info(f"API端点: {self.api_base}")
        logger.info(f"模型: {self.model}")

    def _load_prompt(self) -> str:
        """加载提取 Prompt"""
        prompt_path = PROMPT_DIR / "gaokao_extract_prompt.txt"
        try:
            with open(prompt_path, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            logger.error(f"Prompt 文件未找到: {prompt_path}")
            raise

    def extract_from_docx(self, file_path: str) -> ExtractedContent:
        """
        从 Word 文档提取内容

        Args:
            file_path: docx 文件路径

        Returns:
            ExtractedContent 包含文本、表格、图片
        """
        logger.info(f"[提取] 开始解析文档: {file_path}")

        doc = Document(file_path)
        filename = Path(file_path).stem

        # 从文件路径推断知识点分类
        category = filename  # 文件名就是分类名

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n".join(paragraphs)
        logger.info(f"[提取] 文本段落数: {len(paragraphs)}, 总字符: {len(full_text)}")

        # 提取表格
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_text = self._table_to_text(table)
            tables.append(f"[表格{table_idx + 1}]\n{table_text}")
        logger.info(f"[提取] 表格数量: {len(tables)}")

        # 提取图片
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    images.append(image_data)
                except Exception as e:
                    logger.warning(f"[提取] 图片提取失败: {e}")
        logger.info(f"[提取] 图片数量: {len(images)}")

        # 将表格内容插入文本
        if tables:
            full_text += "\n\n" + "\n\n".join(tables)

        return ExtractedContent(
            text=full_text,
            tables=tables,
            images=images,
            filename=filename,
            category=category
        )

    def _table_to_text(self, table: Table) -> str:
        """将 Word 表格转为文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _wait_if_needed(self):
        """API 频率控制"""
        current_time = time.time()
        time_since_last = current_time - self.last_request_time
        if time_since_last < self.min_request_interval:
            wait_time = self.min_request_interval - time_since_last
            logger.debug(f"[频率控制] 等待 {wait_time:.1f} 秒")
            time.sleep(wait_time)
        self.last_request_time = time.time()

    def extract_questions(
        self,
        content: ExtractedContent,
        chunk_size: int = 8000
    ) -> List[Dict[str, Any]]:
        """
        调用 AI API 提取题目

        Args:
            content: 提取的文档内容
            chunk_size: 每次处理的文本长度（避免超出 token 限制）

        Returns:
            提取的题目列表
        """
        logger.info(f"[API] 开始提取题目: {content.filename}")

        all_questions = []
        text = content.text

        # 如果文本太长，分块处理
        if len(text) > chunk_size:
            chunks = self._split_text(text, chunk_size)
            logger.info(f"[API] 文本过长，分为 {len(chunks)} 块处理")
        else:
            chunks = [text]

        for chunk_idx, chunk in enumerate(chunks):
            logger.info(f"[API] 处理第 {chunk_idx + 1}/{len(chunks)} 块")

            # 构造 prompt
            prompt = self.prompt_template.replace("{category}", content.category)
            prompt = prompt.replace("{filename}", content.filename)
            prompt = prompt.replace("{content}", chunk)

            # 构建消息
            messages = [{"role": "user", "content": prompt}]

            # 如果有图片，添加到消息中（限制前5张）
            if content.images and chunk_idx == 0:
                image_contents = [{"type": "text", "text": prompt}]
                for img_idx, img_bytes in enumerate(content.images[:5]):
                    base64_img = base64.b64encode(img_bytes).decode('utf-8')
                    image_contents.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{base64_img}"}
                    })
                messages = [{"role": "user", "content": image_contents}]

            try:
                self._wait_if_needed()

                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=messages,
                    max_tokens=16384,
                    temperature=0,
                    timeout=120
                )

                response_text = response.choices[0].message.content
                logger.debug(f"[API] 响应长度: {len(response_text)}")

                # 解析 JSON
                questions = self._parse_response(response_text)
                logger.info(f"[API] 第 {chunk_idx + 1} 块提取到 {len(questions)} 道题")

                all_questions.extend(questions)

            except Exception as e:
                logger.error(f"[API] 处理失败: {e}")
                continue

        logger.info(f"[API] 共提取 {len(all_questions)} 道题目")
        return all_questions

    def _split_text(self, text: str, chunk_size: int) -> List[str]:
        """
        智能分割文本，尽量在题目边界分割
        """
        chunks = []
        current_chunk = ""

        # 按段落分割
        paragraphs = text.split("\n")

        for para in paragraphs:
            # 检测是否是新题目开始（数字.[年份）
            is_new_question = False
            if para.strip():
                import re
                if re.match(r'^\d+\s*[\[．.]', para.strip()):
                    is_new_question = True

            # 如果加上这段会超出限制，且这是新题目开始，则分块
            if len(current_chunk) + len(para) > chunk_size and is_new_question and current_chunk:
                chunks.append(current_chunk)
                current_chunk = para + "\n"
            else:
                current_chunk += para + "\n"

        if current_chunk.strip():
            chunks.append(current_chunk)

        return chunks

    def _parse_response(self, response_text: str) -> List[Dict[str, Any]]:
        """解析 API 返回的 JSON"""
        import re

        # 提取 JSON 部分
        text = response_text.strip()

        # 尝试提取 ```json ... ``` 中的内容
        json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
        if json_match:
            text = json_match.group(1).strip()

        try:
            result = json.loads(text)
            if isinstance(result, list):
                return result
            elif isinstance(result, dict) and "questions" in result:
                return result["questions"]
            else:
                logger.warning(f"[解析] 未知的返回格式")
                return []
        except json.JSONDecodeError as e:
            logger.error(f"[解析] JSON 解析失败: {e}")
            logger.debug(f"[解析] 原始文本: {text[:500]}")
            return []

    def process_directory(self, dir_path: str) -> Dict[str, List[Dict]]:
        """
        处理目录下所有 Word 文档

        Args:
            dir_path: 目录路径

        Returns:
            {filename: [questions]}
        """
        results = {}

        # 查找所有 .doc 和 .docx 文件
        path = Path(dir_path)
        doc_files = list(path.glob("**/*.docx")) + list(path.glob("**/*.doc"))

        logger.info(f"[批量] 找到 {len(doc_files)} 个 Word 文件")

        for file_path in doc_files:
            try:
                # .doc 文件需要先转换（暂时跳过）
                if file_path.suffix == '.doc':
                    logger.warning(f"[批量] 跳过 .doc 文件（需要转换）: {file_path.name}")
                    continue

                content = self.extract_from_docx(str(file_path))
                questions = self.extract_questions(content)
                results[file_path.name] = questions

                # 保存中间结果
                self._save_results(file_path.stem, questions)

            except Exception as e:
                logger.error(f"[批量] 处理失败 {file_path.name}: {e}")
                continue

        return results

    def _save_results(self, filename: str, questions: List[Dict]):
        """保存提取结果到 JSON 文件"""
        output_dir = Path("/home/ubuntu/biology-exam-analyzer/uploads/textbooks/extracted")
        output_dir.mkdir(parents=True, exist_ok=True)

        output_path = output_dir / f"{filename}_extracted.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(questions, f, ensure_ascii=False, indent=2)

        logger.info(f"[保存] 结果已保存: {output_path}")


def process_single_file(file_path: str) -> List[Dict]:
    """
    处理单个文件的便捷函数

    使用方法:
        from gaokao_extractor import process_single_file
        questions = process_single_file("/path/to/file.docx")
    """
    import os

    api_key = os.environ.get("DEEPSEEK_API_KEY")
    api_base = os.environ.get("DEEPSEEK_API_BASE")

    if not api_key:
        raise ValueError("请设置 DEEPSEEK_API_KEY 环境变量")

    extractor = GaokaoExtractor(api_key, api_base)
    content = extractor.extract_from_docx(file_path)
    questions = extractor.extract_questions(content)

    return questions


if __name__ == "__main__":
    # 测试单个文件
    import os

    api_key = os.environ.get("DEEPSEEK_API_KEY")
    api_base = os.environ.get("DEEPSEEK_API_BASE")

    if not api_key:
        print("请设置 DEEPSEEK_API_KEY 环境变量")
        exit(1)

    extractor = GaokaoExtractor(api_key, api_base)

    # 测试文件
    test_file = "/home/ubuntu/biology-exam-analyzer/uploads/textbooks/gaokao_zhenti/7.高考真题分类/2023年（含）前高考真题分类/细胞的分子组成.docx"

    if os.path.exists(test_file):
        content = extractor.extract_from_docx(test_file)
        print(f"文本长度: {len(content.text)}")
        print(f"表格数量: {len(content.tables)}")
        print(f"图片数量: {len(content.images)}")

        # 提取题目
        questions = extractor.extract_questions(content)
        print(f"\n提取到 {len(questions)} 道题目")

        # 保存结果
        extractor._save_results(content.filename, questions)

        if questions:
            print("\n第一道题示例:")
            print(json.dumps(questions[0], ensure_ascii=False, indent=2))
    else:
        print(f"测试文件不存在: {test_file}")
