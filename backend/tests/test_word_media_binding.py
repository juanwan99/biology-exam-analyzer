"""Regression tests for Word question media binding."""
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from difficulty_pipeline import DifficultyPipeline
from document_processor import DocumentProcessor
from word_splitter import WordQuestionSplitter


def _tiny_png(path: Path, color: str) -> Path:
    Image.new("RGB", (32, 20), color=color).save(path)
    return path


def _add_blank_image_paragraph(doc: Document, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(image_path))


def _build_media_docx(tmp_path: Path) -> Path:
    image_a = _tiny_png(tmp_path / "a.png", "red")
    image_b = _tiny_png(tmp_path / "b.png", "green")
    image_c = _tiny_png(tmp_path / "c.png", "blue")

    doc = Document()
    doc.add_paragraph("16. 表型正常的材料如下图所示。")
    _add_blank_image_paragraph(doc, image_a)

    doc.add_paragraph("18. 生态关系结果如下表所示。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "生物种类"
    table.cell(0, 1).text = "鱼甲"
    table.cell(1, 0).text = "消化道内食物组成"
    table.cell(1, 1).text = "河虾"
    doc.add_paragraph("装置如下图。")
    _add_blank_image_paragraph(doc, image_b)

    doc.add_paragraph("19. 实验结果如下表所示。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "组别"
    table.cell(0, 1).text = "α-SMA蛋白相对表达量"
    table.cell(1, 0).text = "第一组"
    table.cell(1, 1).text = "1.04"
    doc.add_paragraph("图1")
    _add_blank_image_paragraph(doc, image_c)

    path = tmp_path / "media.docx"
    doc.save(path)
    return path


def _by_id(questions):
    return {question["id"]: question for question in questions}


def test_word_splitter_keeps_images_from_blank_paragraphs(tmp_path):
    docx_path = _build_media_docx(tmp_path)

    questions = _by_id(WordQuestionSplitter().split(str(docx_path))["questions"])

    assert [item["type"] for item in questions[16]["_media_for_ai"]] == ["image"]
    assert [item["type"] for item in questions[18]["_media_for_ai"]] == ["table", "image"]
    assert [item["type"] for item in questions[19]["_media_for_ai"]] == ["table", "image"]


def test_document_processor_matches_structured_content_by_question_boundary(tmp_path):
    docx_path = _build_media_docx(tmp_path)
    questions = WordQuestionSplitter().split(str(docx_path))["questions"]
    elements = DocumentProcessor.extract_word_content(str(docx_path))["elements"]

    DocumentProcessor.match_elements_to_questions(questions, elements)
    by_id = _by_id(questions)

    assert [item["type"] for item in by_id[16]["structured_content"]] == ["image"]

    q18_structured = by_id[18]["structured_content"]
    assert [item["type"] for item in q18_structured] == ["table", "image"]
    assert "生物种类" in q18_structured[0]["markdown"]

    q19_structured = by_id[19]["structured_content"]
    assert [item["type"] for item in q19_structured] == ["table", "image"]
    assert "α-SMA" in q19_structured[0]["markdown"]


def test_word_splitter_marks_missing_media_for_explicit_image_cue(tmp_path):
    doc = Document()
    doc.add_paragraph("18. 装置如下图。")
    path = tmp_path / "missing-image.docx"
    doc.save(path)

    question = WordQuestionSplitter().split(str(path))["questions"][0]

    assert question["media_integrity"]["status"] == "failed"
    assert "image_media_missing" in question["warnings"]


def test_word_splitter_marks_table_render_failure(monkeypatch, tmp_path):
    def fail_render(self, rows_data):
        raise RuntimeError("render boom")

    monkeypatch.setattr(WordQuestionSplitter, "_render_table_as_image", fail_render)

    doc = Document()
    doc.add_paragraph("18. 实验结果如下表所示。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    path = tmp_path / "broken-table-render.docx"
    doc.save(path)

    question = WordQuestionSplitter().split(str(path))["questions"][0]

    assert question["media_integrity"]["status"] == "failed"
    assert "table_media_render_failed" in question["warnings"]
    assert all(item["type"] != "table" for item in question["_media_for_ai"])


@pytest.mark.asyncio
async def test_difficulty_fails_closed_when_media_integrity_failed(monkeypatch):
    async def should_not_call(*args, **kwargs):
        raise AssertionError("difficulty extraction should not run with broken media integrity")

    monkeypatch.setattr("difficulty_pipeline.extract_big_question_features", should_not_call)

    result = await DifficultyPipeline().evaluate_with_refinement(
        {
            "id": 18,
            "content": "18. 装置如下图。",
            "total_score": 11,
            "media_integrity": {"status": "failed", "warnings": ["image_media_missing"]},
        },
        analysis_result={},
    )

    assert result["analysis_failed"] is True
    assert result["failure_reason"] == "media_integrity_failed"
