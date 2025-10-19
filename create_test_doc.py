#!/usr/bin/env python3
"""创建包含表格和图片的测试Word文档"""
from docx import Document
from docx.shared import Inches
from PIL import Image
import io

# 创建文档
doc = Document()

# 添加标题
doc.add_heading('生物测试题目', 0)

# 添加题目1（带表格）
doc.add_heading('题目1：实验结果分析', level=1)
doc.add_paragraph('根据以下实验数据，判断哪个结论是正确的？')

# 添加表格
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# 表头
header_cells = table.rows[0].cells
header_cells[0].text = '实验组'
header_cells[1].text = '温度(℃)'
header_cells[2].text = '酶活性'

# 数据
data = [
    ('对照组', '25', '+++'),
    ('实验组1', '35', '++++'),
    ('实验组2', '45', '++')
]

for i, (group, temp, activity) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = group
    cells[1].text = temp
    cells[2].text = activity

doc.add_paragraph('')
doc.add_paragraph('A. 温度越高酶活性越强')
doc.add_paragraph('B. 35℃时酶活性最强')
doc.add_paragraph('C. 45℃时酶已失活')
doc.add_paragraph('D. 温度对酶活性无影响')

# 添加题目2（带图片）
doc.add_heading('题目2：细胞结构识别', level=1)
doc.add_paragraph('观察下图，判断该细胞类型：')

# 创建一个简单的测试图片
img = Image.new('RGB', (400, 300), color='lightblue')
from PIL import ImageDraw, ImageFont
draw = ImageDraw.Draw(img)
draw.rectangle([50, 50, 350, 250], outline='black', width=3)
draw.ellipse([150, 100, 250, 200], fill='pink', outline='black', width=2)
draw.text((160, 140), '细胞核', fill='black')

# 保存图片到BytesIO
img_bytes = io.BytesIO()
img.save(img_bytes, format='PNG')
img_bytes.seek(0)

# 添加图片到Word
doc.add_picture(img_bytes, width=Inches(3))

doc.add_paragraph('')
doc.add_paragraph('A. 原核细胞')
doc.add_paragraph('B. 真核细胞')
doc.add_paragraph('C. 病毒')
doc.add_paragraph('D. 无法判断')

# 保存文档
output_path = '/app/test_with_table_and_image.docx'
doc.save(output_path)
print(f'✅ 测试文档已创建: {output_path}')
print(f'   - 包含1个表格（4行3列）')
print(f'   - 包含1张图片（400x300）')
